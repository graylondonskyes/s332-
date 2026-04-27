import argparse, json, re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

def slugify(value: str) -> str:
    value = re.sub(r'[^a-z0-9]+', '-', value.lower()).strip('-')
    return value or 'untitled'

def strip_md(md: str):
    out = []
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            out.append('')
            continue
        line = re.sub(r'^#+\s*', '', line)
        line = re.sub(r'^[-*]\s+', '• ', line)
        line = re.sub(r'[`*_>]', '', line)
        out.append(line)
    return out

def parse_json(files, name, fallback=None):
    fallback = fallback or {}
    try:
        return json.loads(files.get(name, '{}'))
    except Exception:
        return fallback

def build_docx(title, subtitle, author, imprint, body_lines, out_path: Path):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].italic = True
    meta = doc.add_paragraph()
    meta.add_run('Author: ').bold = True
    meta.add_run(author)
    meta.add_run('    Imprint: ').bold = True
    meta.add_run(imprint)
    doc.add_paragraph('')
    for line in body_lines:
        if not line:
            doc.add_paragraph('')
        elif line.startswith('• '):
            doc.add_paragraph(line, style='List Bullet')
        elif line.isupper() or (len(line.split()) <= 8 and ':' not in line and line == line.title()):
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)
    doc.save(out_path)

def build_pdf(title, subtitle, author, imprint, body_lines, out_path: Path):
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles['Title'])]
    if subtitle:
        story.append(Paragraph(subtitle, styles['Italic']))
    story.append(Paragraph(f'Author: {author} &nbsp;&nbsp;&nbsp; Imprint: {imprint}', styles['BodyText']))
    story.append(Spacer(1, 12))
    for line in body_lines:
        if not line:
            story.append(Spacer(1, 8))
        elif line.startswith('• '):
            story.append(Paragraph(line, styles['BodyText']))
        elif line == line.title() or line.isupper():
            story.append(Paragraph(line, styles['Heading2']))
        else:
            story.append(Paragraph(line, styles['BodyText']))
    doc = SimpleDocTemplate(str(out_path), pagesize=LETTER, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    doc.build(story)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--workspace', required=True)
    parser.add_argument('--output_dir', required=True)
    args = parser.parse_args()
    workspace = json.loads(Path(args.workspace).read_text())
    files = workspace.get('files', {})
    mode = workspace.get('mode', 'skydocx')
    metadata = parse_json(files, 'metadata.json')
    edition = parse_json(files, 'edition.json')
    channel = parse_json(files, 'channel.json')
    campaign = parse_json(files, 'campaign.json')
    if mode == 'skydocx':
        title = metadata.get('title') or 'Untitled Release'
        subtitle = metadata.get('subtitle') or 'Canonical author manufacturing lane'
        author = metadata.get('author') or 'Operator'
        imprint = metadata.get('imprint') or 'SOLEnterprises'
        slug = metadata.get('slug') or slugify(title)
        body = strip_md(files.get('manuscript.md', ''))
    else:
        title = channel.get('channelTitle') or metadata.get('title') or 'Untitled Editorial'
        subtitle = campaign.get('releaseHook') or 'Synchronized editorial launch lane'
        author = channel.get('author') or metadata.get('author') or 'Operator'
        imprint = 'SkyeBlog'
        slug = channel.get('slug') or metadata.get('slug') or slugify(title)
        body = strip_md(files.get('post.md', ''))
    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f'{slug}-master.docx'
    pdf_path = out_dir / f'{slug}-suite.pdf'
    build_docx(title, subtitle, author, imprint, body, docx_path)
    build_pdf(title, subtitle, author, imprint, body, pdf_path)
    manifest = {
        'schema': 'skye.publishing.binary-job',
        'version': '2.4.0',
        'mode': mode,
        'title': title,
        'slug': slug,
        'edition': edition.get('editionNumber') if isinstance(edition, dict) else None,
        'outputs': [docx_path.name, pdf_path.name]
    }
    manifest_path = out_dir / f'{slug}-binary-job.json'
    manifest_path.write_text(json.dumps(manifest, indent=2) + '\n', encoding='utf-8')
    print(json.dumps({
        'ok': True,
        'mode': mode,
        'slug': slug,
        'docx_path': str(docx_path),
        'pdf_path': str(pdf_path),
        'manifest_path': str(manifest_path)
    }))

if __name__ == '__main__':
    main()
